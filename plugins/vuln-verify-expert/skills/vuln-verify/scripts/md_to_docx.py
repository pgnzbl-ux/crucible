#!/usr/bin/env python
"""
Markdown to Word Document Converter for Vulnerability Reports
Converts report.md to a .docx file with images and proper formatting.
"""

import re
import os
import sys

def check_dependencies():
    """Check if required dependencies are installed."""
    missing = []
    try:
        import docx
    except ImportError:
        missing.append('python-docx')

    if missing:
        print("Missing dependencies:", file=sys.stderr)
        for dep in missing:
            print(f"  - {dep}", file=sys.stderr)
        print("\\nRequired packages are missing from the fixed report-converter image; do not install them interactively.", file=sys.stderr)
        return False
    return True

# Check dependencies before proceeding
if not check_dependencies():
    sys.exit(1)

import docx
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from dataclasses import dataclass
from typing import List


@dataclass
class Run:
    """A styled text fragment produced by the inline parser."""
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


def parse_inline(text: str) -> List[Run]:
    """Parse inline markdown (code/bold/italic) into a flat list of Runs.

    Two-phase: phase 1 splits on backticks (code segments are protected
    from further markdown parsing); phase 2 splits each non-code segment
    on **bold** and *italic*/_italic_. Code segments inherit the
    bold/italic context of the surrounding non-code segment.
    """
    if not text:
        return []

    # Phase 1: split on single backticks. Odd indices are code content.
    parts = text.split('`')
    # Each element: (content, is_code)
    segments = []
    for idx, part in enumerate(parts):
        if part == '':
            continue
        segments.append((part, idx % 2 == 1))

    runs: List[Run] = []
    for content, is_code in segments:
        if is_code:
            runs.append(Run(content, code=True))
        else:
            runs.extend(_parse_emphasis(content))
    return runs


def _parse_emphasis(text: str) -> List[Run]:
    """Split a non-code text segment on **bold** and *italic*/_italic_."""
    if not text:
        return []
    # ** must be matched before single * to avoid conflict.
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)')
    runs: List[Run] = []
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append(Run(text[pos:m.start()]))
        if m.group(2) is not None:          # **bold**
            runs.append(Run(m.group(2), bold=True))
        elif m.group(3) is not None:        # *italic*
            runs.append(Run(m.group(3), italic=True))
        elif m.group(4) is not None:        # _italic_
            runs.append(Run(m.group(4), italic=True))
        pos = m.end()
    if pos < len(text):
        runs.append(Run(text[pos:]))
    return runs


from docx.oxml import OxmlElement


def _set_paragraph_shading(paragraph, fill: str) -> None:
    """Add a background fill (w:shd) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _set_paragraph_border(paragraph, color: str = 'D9D9D9') -> None:
    """Add a single-line border around a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '1')
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)


def _set_run_monospace(run, size_pt: int = 9) -> None:
    """Apply Consolas (latin) + 宋体 (eastAsia) monospace styling to a run."""
    run.font.name = 'Consolas'
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')


def add_code_block(doc, code: str) -> None:
    """Append a code block: monospace font, light-gray shading, border.

    Content is written verbatim — no inline markdown parsing, so
    backticks and asterisks are preserved as literals.
    """
    p = doc.add_paragraph()
    _set_paragraph_shading(p, 'F2F2F2')
    _set_paragraph_border(p, 'D9D9D9')
    run = p.add_run(code)
    _set_run_monospace(run, 9)

def sanitize_filename(name):
    """Remove or replace Windows special characters from filename."""
    # Replace problematic characters with underscore
    replacements = [':', '/', '\\', '<', '>', '|', '?', '*', '"']
    for char in replacements:
        name = name.replace(char, '_')

    # Collapse consecutive underscores
    while '__' in name:
        name = name.replace('__', '_')

    # Remove leading/trailing underscores and spaces
    name = name.strip('_').strip()

    # If empty, use default
    if not name:
        name = "漏洞报告"

    return name

def extract_title(md_content):
    """Extract the first H1 heading from markdown content."""
    match = re.search(r'^#\s+(.+)$', md_content, re.MULTILINE)
    if match:
        return match.group(1).strip()
    return "漏洞报告"

def is_table_separator(line):
    """Check if a line is a markdown table separator."""
    stripped = line.strip()
    # Match patterns like |---|, | :---: |, etc.
    return bool(re.match(r'^\|[\s\-\:|]+\|$', stripped))

def parse_markdown(content):
    """Simple markdown parser to extract sections.

    Images are appended to the content list in document order so they
    render inline where they appear, not deferred to the end.
    """
    lines = content.split('\n')
    sections = []
    current_section = {'level': 0, 'title': '', 'content': []}

    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            # Save previous section
            if current_section['title'] or current_section['content']:
                sections.append(current_section)

            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            current_section = {'level': level, 'title': title, 'content': []}
            i += 1
            continue

        # Image (inline in document order)
        img_match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            current_section['content'].append(('image', (alt_text, img_path)))
            i += 1
            continue

        # Code block
        code_start_match = re.match(r'^```(\w*)$', line.strip())
        if code_start_match:
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            current_section['content'].append(('code', '\n'.join(code_lines)))
            i += 1
            continue

        # Table detection
        if line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            # Collect all lines that start with | (including separator lines)
            while i < len(lines) and (lines[i].strip().startswith('|') or is_table_separator(lines[i])):
                table_lines.append(lines[i])
                i += 1
            current_section['content'].append(('table', table_lines))
            continue

        # Regular content
        if line.strip():
            current_section['content'].append(('text', line))

        i += 1

    # Save last section
    if current_section['title'] or current_section['content']:
        sections.append(current_section)

    return sections

def add_table_from_markdown(doc, table_lines):
    """Convert markdown table to docx table."""
    rows_data = []
    for tl in table_lines:
        # Skip separator lines
        if is_table_separator(tl):
            continue
        cells = [c.strip() for c in tl.strip('|').split('|')]
        if cells and any(c for c in cells):  # Skip empty rows
            rows_data.append(cells)

    if not rows_data or not rows_data[0]:
        return

    # Ensure all rows have the same number of columns
    num_cols = max(len(row) for row in rows_data)
    for row in rows_data:
        while len(row) < num_cols:
            row.append('')

    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if cell_text is None:
                cell_text = ''
            try:
                paragraph = row.cells[j].paragraphs[0]
                paragraph.clear()
                add_runs_from_inline(paragraph, str(cell_text))
            except Exception:
                # Fallback: just set the text
                row.cells[j].text = str(cell_text) if cell_text else ''

def add_runs_from_inline(paragraph, text: str) -> None:
    """Append inline-parsed runs to an existing paragraph.

    Shared by body paragraphs, list items, and table cells so that
    inline code / bold / italic render consistently everywhere.
    """
    for run in parse_inline(text):
        r = paragraph.add_run(run.text)
        if run.bold:
            r.bold = True
        if run.italic:
            r.italic = True
        if run.code:
            _set_run_monospace(r, 10)

def add_formatted_text(doc, text):
    """Add a paragraph, bullet, or numbered list item with inline formatting."""
    if not text:
        return

    text = text.strip()
    if not text:
        return

    # Bullet list
    if text.startswith('- ') or text.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        add_runs_from_inline(p, text[2:])
        return

    # Numbered list
    num_match = re.match(r'^(\d+)\.\s+(.+)$', text)
    if num_match:
        p = doc.add_paragraph(style='List Number')
        add_runs_from_inline(p, num_match.group(2))
        return

    # Regular paragraph
    p = doc.add_paragraph()
    add_runs_from_inline(p, text)

def _resolve_image_path(img_rel_path: str, md_dir: str) -> str:
    """Resolve a markdown image path to an absolute filesystem path.

    Handles Unix/Windows separators and a fallback to an img/ subdir.
    """
    img_rel_path = img_rel_path.replace('/', os.sep).replace('\\', os.sep)
    if os.path.isabs(img_rel_path):
        return img_rel_path if os.path.exists(img_rel_path) else ''
    img_full_path = os.path.join(md_dir, img_rel_path)
    if os.path.exists(img_full_path):
        return img_full_path
    # Fallback: try img/ subdir using only the basename
    alt_path = os.path.join(md_dir, 'img', os.path.basename(img_rel_path))
    if os.path.exists(alt_path):
        return alt_path
    return ''


def add_image_block(doc, alt: str, img_rel_path: str, md_dir: str, fig_counter: list) -> None:
    """Append an inline image (with caption) at the current document position.

    fig_counter is a one-element list used as a mutable counter so figure
    numbers increment across sections in document order.
    """
    img_full_path = _resolve_image_path(img_rel_path, md_dir)
    if not img_full_path:
        doc.add_paragraph(f'[图片缺失: {img_rel_path}]')
        return

    fig_counter[0] += 1
    caption = doc.add_paragraph()
    caption.add_run(f'图 {fig_counter[0]}: {alt}').italic = True

    try:
        doc.add_picture(img_full_path, width=Inches(6))
    except Exception:
        doc.add_paragraph(f'[图片加载失败: {img_full_path}]')
    doc.add_paragraph()  # Spacing

def _section_index(md_content: str) -> dict:
    """Return a {title -> line_no} map of every heading in the markdown.

    Used to validate that the report has all required sections before we burn any
    effort generating a docx that downstream consumers will reject for being
    incomplete.
    """
    index = {}
    for ln_no, line in enumerate(md_content.splitlines(), 1):
        m = re.match(r'^(#{1,6})\s+(.+?)\s*$', line)
        if m:
            title = m.group(2).strip()
            if title not in index:
                index[title] = ln_no
    return index


REQUIRED_SECTIONS = [
    '1. 产品介绍',
    '2. 漏洞描述',
    '3. 影响范围',
    '4. 漏洞详情',
    '5. 漏洞复现',
    '6. POC',
    '7. 修复建议',
    '8. 报送判定',
]


def _normalize_heading(title: str) -> str:
    """Strip parentheses / trailing decoration for fuzzy matching.

    `## 8. 报送判定（文字反馈）` and `## 8. 报送判定` should both match
    `## 8. 报送判定`. Likewise `## 4. 漏洞详情（X）` ↔ `## 4. 漏洞详情`.
    """
    t = title.strip()
    # Strip a trailing parenthetical, e.g. "报送判定（文字反馈）" → "报送判定".
    paren = re.search(r'[（(][^()）]*[)）]\s*$', t)
    if paren:
        t = t[:paren.start()].rstrip()
    return t


def _validate_report(md_path: str, md_content: str) -> list:
    """Hard validation that runs BEFORE doc generation. Returns a list of
    human-readable error strings; an empty list means OK.

    Checks:
    1. All required numbered sections (1-8) are present (fuzzy match
       tolerates a trailing parenthetical like "（文字反馈）").
    2. §5.1 explicitly mentions the transport / connector shape. The
       keywords are transport-agnostic (HTTPS / HTTP / WebSocket / gRPC /
       MQTT), so HTTP-only or other-transport reports are equally accepted.
    3. Each "步骤 N" inside §5.2 has a paired `![<alt>](img/stepN_xxx.png)`
       image reference immediately following it. The step marker tolerates
       `: description text` after the bold-closed number.
    4. Every image reference resolves to an actual file under img/.
    5. No image syntax inside a markdown table cell — the table renderer
       writes cells as plain text, so `![](img/x.png)` inside a table would
       silently appear as literal text (images dropped from the docx).

    When any check fails, return errors and caller's exit code reflects it.
    """
    errors = []
    md_dir = os.path.dirname(os.path.abspath(md_path))

    # Check 5: images must NOT be embedded inside markdown table cells.
    for ln_no, line in enumerate(md_content.splitlines(), 1):
        if line.lstrip().startswith('|') and re.search(r'!\[[^\]]*\]\([^\)]+\)', line):
            errors.append(
                f'[SCREENSHOT-IN-TABLE] line {ln_no}: image syntax inside a '
                f'table cell is not rendered by the docx converter (cells are '
                f'plain text). Move the image to its own standalone line: '
                f'![<alt>](img/stepN_xxx.png)'
            )

    sections = _section_index(md_content)
    normalized = {_normalize_heading(t): ln for t, ln in sections.items()}
    for required in REQUIRED_SECTIONS:
        if required not in normalized:
            errors.append(
                f'[SECTION] missing required section "{required}" — '
                f'report_template.md mandates all 8 numbered sections.'
            )

    section_51_start = sections.get('5.1 环境准备') or normalized.get('5.1 环境准备')
    section_52_start = sections.get('5.2 复现步骤') or normalized.get('5.2 复现步骤')

    # Check 2: §5.1 mentions at least one transport / connector keyword.
    # Keywords chosen to be transport-agnostic:
    #   connector / 连接器 / 端口 / TLS / HTTP / 协议 / 通道 / WebSocket / gRPC / MQTT
    # So HTTP-only or WebSocket reports are equally accepted.
    if section_51_start:
        lines = md_content.splitlines()
        section_51_end = (section_52_start or len(lines) + 1) - 1
        section_51_idx = section_51_start - 1
        section_51_text = '\n'.join(lines[section_51_idx:section_51_end])
        transport_keywords = [
            'Connector', '连接器', '端口',
            'TLS', 'HTTPS', 'HTTP', 'websocket', 'WebSocket',
            'gRPC', 'grpc', 'MQTT', 'AMQP', 'QUIC',
            '协议', 'channel', '通道', '终结', 'terminat',
        ]
        if not any(kw in section_51_text for kw in transport_keywords):
            errors.append(
                '[DEPLOYMENT-SURFACE] §5.1 must explicitly describe the '
                'transport / connector shape. Mention at least one of: '
                'Connector / 连接器 / 端口 / TLS / HTTPS / WebSocket / gRPC / '
                'MQTT / 协议 / 通道. See references/deployment-surface.md.'
            )

    # Check 3: each "**步骤 N**" in §5.2 must have an image somewhere between
    # itself and the next "**步骤 N+1**" (or end of §5.2). Tolerant of
    # intervening prose, code blocks, bullets — the report author usually
    # writes one paragraph of context before dropping the screenshot.
    if section_52_start:
        lines = md_content.splitlines()
        next_required_section = (
            sections.get('5.3 结果验证') or
            normalized.get('5.3 结果验证') or
            len(lines) + 1
        )
        next_required_idx = next_required_section - 1
        section_52_idx = section_52_start - 1
        section_52_lines = lines[section_52_idx:next_required_idx]

        # Find line indices of each step heading inside §5.2.
        step_re = re.compile(r'^\s*\*\*步骤\s+(\d+)\*\*\s*[:：]?')
        img_re = re.compile(
            r'!\[[^\]]*\]\(\s*(img/[^\s)]+\.png)\s*\)', re.IGNORECASE
        )

        step_line_indices = []  # list of (step_n, line_idx_within_section_52)
        for idx, line in enumerate(section_52_lines):
            sm = step_re.match(line)
            if sm:
                step_line_indices.append((int(sm.group(1)), idx))

        # For each step (except the last), search its window
        # [this_idx, next_step_idx). For the last step, [this_idx, end).
        for i, (step_n, this_idx) in enumerate(step_line_indices):
            window_end = (
                step_line_indices[i + 1][1] if i + 1 < len(step_line_indices)
                else len(section_52_lines)
            )
            block = section_52_lines[this_idx:window_end]
            found = None
            for ln in block:
                im = img_re.search(ln)
                if im:
                    found = im.group(1)
                    break
            if not found:
                errors.append(
                    f'[SCREENSHOT-MISSING] "**步骤 {step_n}**" in §5.2 has '
                    f'no inline image between this step and the next step '
                    f'(or end of §5.2). Drop a real screenshot anywhere in '
                    f'this block: ![<alt>](img/step{step_n}_xxx.png).'
                )
                continue
            img_rel = found
            resolved = _resolve_image_path(img_rel, md_dir)
            if not resolved:
                errors.append(
                    f'[SCREENSHOT-FILE-MISSING] step {step_n} references '
                    f'"{img_rel}" but no file exists at that path under the '
                    f'report directory. Either save the screenshot or remove '
                    f'the broken reference.'
                )

    return errors


def convert_md_to_docx(md_path, output_dir=None):
    """Convert markdown report to Word document."""

    # Resolve full path
    md_path = os.path.abspath(md_path)

    # Read markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Pre-flight validation — fail loudly on missing sections or screenshots.
    preflight_errors = _validate_report(md_path, md_content)
    if preflight_errors:
        print('Refusing to generate docx: report.md failed preflight checks:',
              file=sys.stderr)
        for e in preflight_errors:
            print(f'  - {e}', file=sys.stderr)
        print('\nFix the markdown, then rerun.', file=sys.stderr)
        sys.exit(2)

    # Extract title and sanitize for filename
    title = extract_title(md_content)
    safe_title = sanitize_filename(title)

    # Determine output path
    if output_dir is None:
        output_dir = os.path.dirname(md_path)

    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    docx_filename = safe_title + '.docx'
    docx_path = os.path.join(output_dir, docx_filename)

    # Create document
    doc = docx.Document()

    # Set default font for Chinese
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Add title
    title_heading = doc.add_heading(title, 0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get image directory (same as md file's directory)
    md_dir = os.path.dirname(md_path)

    # Parse markdown
    sections = parse_markdown(md_content)

    fig_counter = [0]  # mutable counter so figure numbers span sections in order
    for section in sections:
        # Add heading
        if section['title']:
            level = min(section['level'], 9)  # docx only supports 1-9
            if level == 0:
                level = 1
            doc.add_heading(section['title'], level)

        # Add content in document order (text/code/table/image interleaved)
        for content_type, content in section['content']:
            if content_type == 'text':
                add_formatted_text(doc, content)
            elif content_type == 'code':
                add_code_block(doc, content)
            elif content_type == 'table':
                add_table_from_markdown(doc, content)
            elif content_type == 'image':
                alt, img_rel_path = content
                add_image_block(doc, alt, img_rel_path, md_dir, fig_counter)

    # Save document
    try:
        doc.save(docx_path)
        return docx_path
    except Exception as e:
        print(f"Error saving document: {e}", file=sys.stderr)
        raise

if __name__ == '__main__':
    # Determine script's own directory for resource location
    if getattr(sys, 'frozen', False):
        # Running as compiled executable
        script_dir = os.path.dirname(sys.executable)
    else:
        script_dir = os.path.dirname(os.path.abspath(__file__))

    if len(sys.argv) < 2:
        print("Usage: md_to_docx.py <report.md> [output_dir]")
        print(f"  Script location: {script_dir}")
        sys.exit(1)

    md_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        output_path = convert_md_to_docx(md_path, output_dir)
        print(f"Document saved to: {output_path}")
    except SystemExit:
        raise
    except Exception as e:
        print(f"Failed to convert: {e}", file=sys.stderr)
        sys.exit(1)
