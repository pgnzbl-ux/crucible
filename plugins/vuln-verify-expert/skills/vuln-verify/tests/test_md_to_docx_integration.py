import unittest
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))
import docx
from docx.oxml.ns import qn
from md_to_docx import add_code_block, add_formatted_text, add_table_from_markdown, convert_md_to_docx


def _make_minimal_png(path):
    """Write a valid 1x1 PNG using only the stdlib (no PIL)."""
    import zlib, struct
    width = height = 1
    raw = b'\x00' + b'\xff\x00\x00'  # filter byte + one RGB pixel (red)
    compressed = zlib.compress(raw)

    def chunk(tag, data):
        return struct.pack('>I', len(data)) + tag + data + struct.pack('>I', zlib.crc32(tag + data) & 0xffffffff)

    sig = b'\x89PNG\r\n\x1a\n'
    ihdr = struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0)  # 8-bit, color type 2 (RGB)
    with open(path, 'wb') as f:
        f.write(sig)
        f.write(chunk(b'IHDR', ihdr))
        f.write(chunk(b'IDAT', compressed))
        f.write(chunk(b'IEND', b''))


class TestCodeBlock(unittest.TestCase):
    def _shading_fill(self, paragraph):
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is None:
            return None
        shd = pPr.find(qn('w:shd'))
        if shd is None:
            return None
        return shd.get(qn('w:fill'))

    def test_code_block_has_shading_and_consolas(self):
        doc = docx.Document()
        add_code_block(doc, "const x = `template ${1}`")
        p = doc.paragraphs[-1]
        # 浅灰背景
        self.assertEqual(self._shading_fill(p), "F2F2F2")
        # 等宽字体
        run = p.runs[0]
        self.assertEqual(run.font.name, "Consolas")
        # 内容原样，反引号未被解析
        self.assertIn("`template", run.text)

    def test_code_block_preserves_asterisks(self):
        doc = docx.Document()
        add_code_block(doc, "**not bold**")
        self.assertEqual(doc.paragraphs[-1].runs[0].text, "**not bold**")
        self.assertFalse(doc.paragraphs[-1].runs[0].bold)


class TestFormattedText(unittest.TestCase):
    def test_inline_code_in_paragraph(self):
        doc = docx.Document()
        add_formatted_text(doc, "设置 `securityLevel`")
        p = doc.paragraphs[-1]
        runs = p.runs
        self.assertTrue(any(r.font.name == "Consolas" for r in runs))
        code_run = next(r for r in runs if r.font.name == "Consolas")
        self.assertEqual(code_run.text, "securityLevel")

    def test_bold_in_paragraph(self):
        doc = docx.Document()
        add_formatted_text(doc, "**加粗**")
        self.assertTrue(doc.paragraphs[-1].runs[0].bold)

    def test_bullet_list_with_code(self):
        doc = docx.Document()
        add_formatted_text(doc, "- 项 `code`")
        p = doc.paragraphs[-1]
        self.assertIn("List", p.style.name)
        self.assertTrue(any(r.font.name == "Consolas" for r in p.runs))


class TestTableInlineFormat(unittest.TestCase):
    def test_cell_with_inline_code(self):
        doc = docx.Document()
        table_lines = [
            "| 文件 | 行号 |",
            "|------|------|",
            "| `mermaid.tsx` | 35 |",
        ]
        add_table_from_markdown(doc, table_lines)
        table = doc.tables[-1]
        cell = table.rows[1].cells[0]
        self.assertTrue(any(r.font.name == "Consolas" for r in cell.paragraphs[0].runs))
        self.assertEqual(cell.paragraphs[0].runs[0].text, "mermaid.tsx")

    def test_cell_with_bold(self):
        doc = docx.Document()
        add_table_from_markdown(doc, ["| a | b |", "|---|---|", "| **x** | y |"])
        cell = doc.tables[-1].rows[1].cells[0]
        self.assertTrue(cell.paragraphs[0].runs[0].bold)


class TestImageOrder(unittest.TestCase):
    """An image must render in document order where it appears in markdown,
    not be deferred to the end of its section."""

    def _body_block_kinds(self, doc):
        """Walk top-level body elements in order; classify each as
        'text:<first-run-text>' or 'image'."""
        body = doc.element.body
        kinds = []
        for child in body.iterchildren():
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                # paragraph with an inline drawing = image
                drawings = child.findall('.//' + qn('w:drawing'))
                if drawings:
                    kinds.append('image')
                else:
                    texts = [t.text for t in child.findall('.//' + qn('w:t')) if t.text]
                    kinds.append('text:' + ''.join(texts))
            elif tag == 'tbl':
                kinds.append('table')
        return kinds

    def test_image_appears_between_paragraphs(self):
        import tempfile, shutil
        tmp = tempfile.mkdtemp()
        try:
            img = os.path.join(tmp, 'img', 'dot.png')
            os.makedirs(os.path.dirname(img))
            _make_minimal_png(img)
            md = os.path.join(tmp, 'r.md')
            with open(md, 'w', encoding='utf-8') as f:
                f.write("# T\n\n步骤1之前文字\n\n![pic](img/dot.png)\n\n步骤1之后文字\n")
            out = convert_md_to_docx(md, tmp)
            doc = docx.Document(out)
            kinds = self._body_block_kinds(doc)
            img_idx = kinds.index('image')
            before_idx = next(i for i, k in enumerate(kinds) if k.startswith('text:步骤1之前'))
            after_idx = next(i for i, k in enumerate(kinds) if k.startswith('text:步骤1之后'))
            self.assertLess(before_idx, img_idx, '文字A 必须在图片之前')
            self.assertLess(img_idx, after_idx, '图片必须在文字B之前')
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
